from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


OUT = Path("output/docx/milad_hospital_translation.docx")
ASSET_DIR = Path("tmp/docx_assets")


PAGES = [
    {
        "patient": "Ms. Zohreh Khanom Eftekhari",
        "father": "Mirza Ali Akbar",
        "diagnosis": "sepsis and renal failure",
        "start": "1401/06/20",
        "end": "1401/06/28",
        "ward": "ICU",
        "doctor": "Dr. Saba",
    },
    {
        "patient": "Ms. Zohreh Khanom Eftekhari",
        "father": "Mirza Ali Akbar",
        "diagnosis": "sepsis",
        "start": "1401/01/21",
        "end": "1401/01/25",
        "ward": "ICU",
        "doctor": "Dr. Saba",
    },
    {
        "patient": "Ms. Zohreh Khanom Eftekhari",
        "father": "Mirza Ali Akbar",
        "diagnosis": "sepsis",
        "start": "1401/06/13",
        "end": "1401/06/19",
        "ward": "ICU",
        "doctor": "Dr. Saba",
    },
]


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_page_border(section):
    sect_pr = section._sectPr
    pg_borders = sect_pr.find(qn("w:pgBorders"))
    if pg_borders is None:
        pg_borders = OxmlElement("w:pgBorders")
        sect_pr.append(pg_borders)
    pg_borders.set(qn("w:offsetFrom"), "page")
    for side in ("top", "left", "bottom", "right"):
        border = pg_borders.find(qn(f"w:{side}"))
        if border is None:
            border = OxmlElement(f"w:{side}")
            pg_borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "8")
        border.set(qn("w:space"), "24")
        border.set(qn("w:color"), "000000")


def add_run(paragraph, text, bold=False, size=11, color=None, font="Calibri"):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return run


def blank(paragraph, size=11):
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1
    add_run(paragraph, "", size=size)


def make_stamp_image(path):
    img = Image.new("RGBA", (560, 360), (255, 255, 255, 0))
    draw = ImageDraw.Draw(img)
    blue = (45, 119, 200, 255)
    try:
        font_bold = ImageFont.truetype("Arial Bold.ttf", 28)
        font = ImageFont.truetype("Arial.ttf", 24)
        font_small = ImageFont.truetype("Arial.ttf", 20)
    except OSError:
        font_bold = ImageFont.load_default()
        font = ImageFont.load_default()
        font_small = ImageFont.load_default()
    draw.ellipse((80, 25, 390, 335), outline=blue, width=7)
    draw.text((205, 145), "STAMP", fill=blue, font=font_bold, anchor="mm")
    draw.text((20, 280), "Health Information Management Department", fill=(0, 0, 0, 255), font=font)
    draw.text((20, 325), "Milad Hospital Kashan", fill=(0, 0, 0, 255), font=font)
    # faint ring text placeholder, only to visually echo the original seal.
    draw.text((115, 45), "Milad Hospital Kashan", fill=(45, 119, 200, 120), font=font_small)
    img.save(path)


def make_signature_image(path):
    img = Image.new("RGBA", (420, 110), (255, 255, 255, 0))
    draw = ImageDraw.Draw(img)
    blue = (32, 58, 122, 255)
    draw.line((20, 35, 390, 10), fill=blue, width=4)
    try:
        font = ImageFont.truetype("Arial.ttf", 24)
    except OSError:
        font = ImageFont.load_default()
    draw.text((45, 70), "Signature", fill=(0, 0, 0, 255), font=font)
    img.save(path)


def apply_document_defaults(doc):
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    section.header_distance = Cm(1.0)
    section.footer_distance = Cm(1.0)
    set_page_border(section)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10


def add_top_line(doc):
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Inches(3.25)
    table.columns[1].width = Inches(3.25)
    for cell in table.rows[0].cells:
        set_cell_margins(cell, top=0, bottom=0)
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.space_after = Pt(0)
    left = table.cell(0, 0).paragraphs[0]
    add_run(left, "RM-FO-03", bold=True, size=12, font="Times New Roman")
    right = table.cell(0, 1).paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(right, "English Translation", size=11)


def add_certificate_page(doc, data, is_last=False):
    add_top_line(doc)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(52)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "Milad Hospital", bold=True, size=26, font="Times New Roman")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(82)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "In the Name of God", bold=True, size=14, font="Times New Roman")

    body = (
        f"Respectfully, this is to certify that {data['patient']}, "
        f"daughter of {data['father']}, with the diagnosis of {data['diagnosis']}, "
        f"was hospitalized from {data['start']} to {data['end']} in the {data['ward']} "
        f"ward of Milad Hospital. The treating physician was {data['doctor']}. "
        "This certificate has been issued at the request of the above-named person "
        "for presentation and has no other validity."
    )
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(130)
    p.paragraph_format.line_spacing = 1.35
    p.paragraph_format.space_after = Pt(0)
    add_run(p, body, size=13, font="Times New Roman")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(130)
    p.paragraph_format.space_after = Pt(0)
    blank(p)

    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Inches(3.25)
    table.columns[1].width = Inches(3.25)
    stamp_cell, sig_cell = table.rows[0].cells
    for cell in (stamp_cell, sig_cell):
        set_cell_margins(cell, top=0, bottom=0)
    stamp_p = stamp_cell.paragraphs[0]
    stamp_p.paragraph_format.space_after = Pt(0)
    stamp_p.add_run().add_picture(str(ASSET_DIR / "stamp.png"), width=Inches(2.95))
    sig_p = sig_cell.paragraphs[0]
    sig_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sig_p.paragraph_format.space_after = Pt(0)
    sig_p.add_run().add_picture(str(ASSET_DIR / "signature.png"), width=Inches(2.6))

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(16)
    note.paragraph_format.space_after = Pt(0)
    add_run(
        note,
        "Stamp text translated according to the visible seal: "
        "Health Information Management Department / Milad Hospital Kashan",
        size=9,
        font="Times New Roman",
    )

    if not is_last:
        doc.add_section(WD_SECTION.NEW_PAGE)
        set_page_border(doc.sections[-1])


def main():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    make_stamp_image(ASSET_DIR / "stamp.png")
    make_signature_image(ASSET_DIR / "signature.png")

    doc = Document()
    apply_document_defaults(doc)
    doc.core_properties.title = "Milad Hospital English Translation"
    doc.core_properties.author = ""

    for idx, page in enumerate(PAGES):
        add_certificate_page(doc, page, is_last=idx == len(PAGES) - 1)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
